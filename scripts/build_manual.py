# -*- coding: utf-8 -*-
"""
Build the Medical-Supply ERP User Manual (English + Arabic) from a single
content source into professionally formatted Word (.docx) and HTML files.
The HTML is converted to PDF with the Odoo container's wkhtmltopdf.

Outputs (docs/manual/):
    Medical-Supply_ERP_User_Manual_EN.docx / _EN.html
    Medical-Supply_ERP_User_Manual_AR.docx / _AR.html   (right-to-left)

Usage:  python3 scripts/build_manual.py
"""
import html as _html
import json
import os

BRAND = "0E6E8E"        # teal
BRAND_DARK = "094E66"
ACCENT = "1F7A4D"       # medical green
LIGHT = "EAF3F6"
GREY = "F4F6F7"

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, "docs", "manual")
FONT_DIR = os.path.join(ROOT, "scripts", "fonts")

META = {
    "en": dict(title="Medical-Supply ERP", subtitle="User Manual",
               org="Sudan MedSupply Co.", platform="Built on Odoo 18 Community",
               version="Version 3.0", date="June 2026", toc="Table of Contents",
               figures="List of Figures",
               labels={"tip": "TIP", "warn": "IMPORTANT", "note": "NOTE", "fig": "Figure"}),
    "ar": dict(title="نظام إدارة موارد المستلزمات الطبية", subtitle="دليل المستخدم",
               org="شركة السودان للمستلزمات الطبية", platform="مبني على أودو 18 المجتمعي",
               version="الإصدار 3.0", date="يونيو 2026", toc="جدول المحتويات",
               figures="قائمة الأشكال",
               labels={"tip": "نصيحة", "warn": "هام", "note": "ملاحظة", "fig": "شكل"}),
}

# ----------------------------------------------------------------------------
class Content:
    """Collects manual blocks. Inline **bold** is supported in text."""
    def __init__(self):
        self.C = []
    def h1(self, t): self.C.append(("h1", t))
    def h2(self, t): self.C.append(("h2", t))
    def h3(self, t): self.C.append(("h3", t))
    def p(self, t): self.C.append(("p", t))
    def ul(self, items): self.C.append(("ul", items))
    def ol(self, items): self.C.append(("ol", items))
    def table(self, headers, rows): self.C.append(("table", headers, rows))
    def tip(self, title, t): self.C.append(("tip", title, t))
    def warn(self, title, t): self.C.append(("warn", title, t))
    def note(self, title, t): self.C.append(("note", title, t))
    def fig(self, filename, caption): self.C.append(("fig", filename, caption))

# ====================================================== JSON CONTENT LOADER
# Chapters produced by scripts/wf_manual_content.js live in docs/manual/_content/
# as <key>.json with {title_en,title_ar,blocks_en,blocks_ar}. This is the single
# source of truth for the manual; fill_en/fill_ar below are kept only as fallback.
CHAPTER_ORDER = [
    "intro", "interface", "roles", "config_company", "config_partners",
    "config_banks", "config_inventory", "products", "warehouse", "procurement",
    "sales", "accounting", "multicurrency", "medical", "journeys",
    "demo_reference", "admin",
]

def content_from_json(lang):
    """Assemble the ordered block-tuple list for a language from the per-chapter
    JSON files; returns [] if no chapter files are present."""
    cdir = os.path.join(OUT_DIR, "_content")
    blocks = []
    found = False
    for key in CHAPTER_ORDER:
        path = os.path.join(cdir, key + ".json")
        if not os.path.exists(path):
            continue
        found = True
        with open(path, encoding="utf-8") as fh:
            data = json.load(fh)
        title = data.get("title_%s" % lang) or data.get("title_en") or key
        blocks.append(("h1", title))
        for b in data.get("blocks_%s" % lang, []):
            t = b.get("t")
            if t in ("h2", "h3", "p"):
                blocks.append((t, b.get("text", "")))
            elif t in ("ul", "ol"):
                blocks.append((t, list(b.get("items", []))))
            elif t == "table":
                blocks.append(("table", list(b.get("headers", [])),
                               [list(r) for r in b.get("rows", [])]))
            elif t in ("tip", "warn", "note"):
                blocks.append((t, b.get("title", ""), b.get("text", "")))
            elif t == "fig":
                blocks.append(("fig", b.get("file", ""), b.get("caption", "")))
    return blocks if found else []

# ============================================================ ENGLISH CONTENT
def fill_en(d):
    d.h1("1. Introduction")
    d.p("This manual describes the **Medical-Supply ERP** — an integrated business system "
        "for a company that imports, stores, and distributes medical supplies and "
        "pharmaceuticals. The system covers the full operational chain: master-data "
        "configuration, inventory and warehousing, procurement, sales, and accounting, with "
        "native **multi-currency** support (Sudanese Pound as the base currency and US Dollar "
        "as a reference currency).")
    d.p("The system is built on **Odoo 18 Community** and uses only Community and open-source "
        "(OCA / Odoo-Mates) modules — no proprietary Enterprise licences are required.")
    d.h2("1.1 Who this manual is for")
    d.ul(["**Administrators** who configure master data — warehouses, items, partners, banks, "
          "currencies and exchange rates.",
          "**Procurement officers** who raise purchase orders and receive goods.",
          "**Sales staff** who quote, sell, and deliver to customers.",
          "**Accountants** who manage invoices, payments, bank reconciliation and reports."])
    d.h2("1.2 Key capabilities at a glance")
    d.table(["Area", "What you can do"],
            [["Configuration", "Define company, currencies & rates, customers, suppliers, banks, cash safes, units, warehouses and product categories."],
             ["Inventory", "Register items with lot/batch and expiry tracking; manage on-hand stock and valuation."],
             ["Warehouse", "Run multiple warehouses and internal locations (e.g. cold storage, quarantine); receipts, deliveries and internal transfers."],
             ["Procurement", "RFQ → Purchase Order → Receipt → Vendor Bill, with vendor price lists and multi-currency purchasing."],
             ["Sales", "Quotation → Sales Order → Delivery → Customer Invoice, with price lists per currency."],
             ["Accounting", "Customer invoices, vendor bills, payments, bank reconciliation, financial reports, fiscal year and dunning."],
             ["Multi-currency", "SDG base + USD reference; dated exchange-rate history that is easy to update."]])

    d.h1("2. System Overview")
    d.h2("2.1 Building blocks")
    d.p("The ERP is organised into cooperating applications. Each application adds menus, "
        "documents and reports, but they share the same products, partners and accounting.")
    d.table(["Application", "Purpose"],
            [["Inventory", "Items, stock, warehouses, locations, lots, expiry, transfers."],
             ["Purchase", "Supplier orders and goods receipts."],
             ["Sales", "Customer quotations, orders and deliveries."],
             ["Invoicing / Accounting", "Invoices, bills, payments, journals, reports."],
             ["Contacts", "Customers and suppliers (a single contact can be both)."],
             ["Settings", "Company, currencies, users, and feature switches."]])
    d.h2("2.2 Logging in")
    d.ol(["Open the system URL in a web browser (for the local demo: **http://localhost:8069**).",
          "If a database list appears, choose **erpmedsupply**.",
          "Enter your user name and password. The default administrator is **admin / admin** "
          "on the demo database — change this password before going live."])
    d.fig("login.png", "The sign-in screen — choose the **erpmedsupply** database and log in.")
    d.warn("Change default credentials", "The demo administrator password (**admin**) and the "
           "database master password must be changed before the system is used with real data.")
    d.h2("2.3 Finding your way around")
    d.ul(["The **top menu bar** switches between applications (Inventory, Purchase, Sales, Accounting…).",
          "Each application has its own **Configuration** sub-menu for setup tasks.",
          "Most documents follow a **draft → confirmed → done** lifecycle shown by status "
          "buttons in the top-right."])
    d.fig("apps_home.png", "The applications home screen, showing the installed business apps.")

    d.h1("3. Configuration (Central Admin)")
    d.p("Configuration is the heart of the system. Complete these steps **in order** before "
        "recording day-to-day transactions, because some settings (especially the base "
        "currency and chart of accounts) are difficult to change once transactions exist.")
    d.h2("3.1 Company")
    d.ol(["Go to **Settings → Users & Companies → Companies** and open your company.",
          "Set the legal name, address, tax number and logo.",
          "Set the **Country** to Sudan."])
    d.h2("3.2 Currencies and exchange rates")
    d.p("The system uses **SDG (Sudanese Pound)** as the company base currency and **USD** as "
        "a reference currency for imports and price quotations.")
    d.ol(["Enable multi-currency: **Settings → Accounting → Currencies → tick Multi-Currencies**.",
          "Activate the currencies you use: **Settings → Currencies**, switch on **SDG** and **USD**.",
          "Confirm the company base currency is **SDG**."])
    d.h3("Recording a new USD exchange rate")
    d.p("Exchange rates are stored as **dated records**, so the full history is preserved and "
        "any past document is valued at the rate that applied on its date.")
    d.ol(["Open **Accounting → Configuration → Currencies** and click **USD**.",
          "Go to the **Rates** tab and click **Add a line**.",
          "Enter the **date** and the rate — type the intuitive figure: **how many SDG equal "
          "1 USD** (e.g. 700).",
          "Save. The new rate applies to all documents dated on or after that day."])
    d.tip("Easy rate updates", "Because each change is a new dated line, you simply add a line "
          "whenever the rate moves — you never overwrite history. Reports and documents "
          "automatically use the correct rate for their date.")
    d.fig("currency_rates.png", "The USD currency record with its dated rate history "
          "(1 USD = 600 → 700 SDG). Type the rate into the **SDG per Unit** column.")
    d.note("Automatic rates (optional)", "Community Odoo does not fetch rates automatically. If "
           "you want scheduled updates, the OCA **currency_rate_update** module can be added "
           "later; otherwise rates are maintained by hand as above.")
    d.h2("3.3 Customers and suppliers")
    d.p("Customers and suppliers are both **Contacts**. A single contact can be both a customer "
        "and a supplier — do not create duplicates.")
    d.ol(["Open **Contacts → New**.",
          "Choose **Company** or **Individual** and enter name, address and phone.",
          "On the **Sales & Purchasing** tab set the price list, payment terms and (for "
          "foreign suppliers) the purchase currency.",
          "On the **Accounting** tab set receivable/payable accounts and add bank accounts."])
    d.table(["Field", "Used for"],
            [["Customer / Vendor", "Determines whether the contact appears in Sales or Purchase."],
             ["Purchase currency", "Default currency for that supplier's orders (e.g. USD for imports)."],
             ["Price list", "Which selling prices and currency apply to this customer."],
             ["Bank accounts", "Needed to register or batch payments."]])
    d.fig("contacts.png", "The Contacts directory — customers and suppliers share one list.")
    d.h2("3.4 Banks and money safes")
    d.p("Money flows are tracked through **journals**. A **bank** needs both a bank record and "
        "a bank journal; a **money safe / cash box** is a **cash journal**.")
    d.ol(["Create the institution: **Accounting → Configuration → Banks → New** (name, SWIFT/BIC).",
          "Create a **Bank** journal: Type = **Bank**. Set a currency for a foreign-currency "
          "account (e.g. a USD bank account).",
          "Create a **Cash** journal for each physical safe: Type = **Cash** (e.g. *Main Cash "
          "Safe (SDG)*, *USD Cash Safe*)."])
    d.note("Why a cash journal is a 'safe'", "There is no separate 'safe' object. Each physical "
           "cash box or till is represented by its own **Cash journal**, so its balance and "
           "movements are tracked and reconciled independently.")
    d.h2("3.5 Units of measure")
    d.p("Enable **Units of Measure** to buy, stock and sell in different units — for example "
        "purchase a *Box of 100* but issue *units*. Define units and conversion factors under "
        "**Inventory → Configuration → Units of Measure**.")
    d.h2("3.6 Warehouses and locations")
    d.ol(["Go to **Inventory → Configuration → Warehouses** and create one per physical site "
          "(name + short code, e.g. *Khartoum Central / KRT*).",
          "Set the number of receipt and delivery steps on each warehouse as needed.",
          "Create internal **Locations**: **Cold Storage (2-8°C)**, **Quarantine**, "
          "**Expired / Damaged**."])
    d.tip("Cold chain", "Temperature-controlled storage is modelled as an internal "
          "**location**. A **putaway rule** can automatically route a product (e.g. insulin) "
          "into Cold Storage when it is received.")
    d.fig("locations.png", "Internal locations under Khartoum Central — including "
          "**Cold Storage (2-8°C)**, **Quarantine** and **Expired / Damaged**.")
    d.h2("3.7 Product categories")
    d.p("Product categories control how stock is costed and removed.")
    d.table(["Setting", "Recommended value", "Effect"],
            [["Costing Method", "FIFO", "Stock is valued first-in-first-out."],
             ["Inventory Valuation", "Manual / Periodic*", "Avoids needing valuation accounts up front."],
             ["Force Removal Strategy", "FEFO", "Goods nearest to expiry are issued first."]])
    d.note("Automated valuation", "*The demo uses periodic valuation so goods movements post "
           "without extra account setup. To post stock value in real time, set the category's "
           "Stock Input/Output/Valuation accounts and a Stock journal, then switch Valuation "
           "to **Automated**.")

    d.h1("4. Inventory & Items")
    d.h2("4.1 Registering an item")
    d.ol(["Open **Inventory → Products → Products → New**.",
          "Enter the name, **Internal Reference** (SKU) and **Barcode**.",
          "Set **Product Type = Goods** and switch on **Track Inventory**.",
          "Assign the **Product Category** (e.g. Pharmaceuticals).",
          "Set the **Cost** and **Sales Price**.",
          "On the **Inventory** tab set **Tracking = By Lots** and enable **Expiration Date** "
          "for medicines and consumables."])
    d.fig("products_list.png", "The medical product catalogue with on-hand quantities.")
    d.fig("product_insulin.png", "Product form for **Insulin Vial 10ml** — tracked by lot, "
          "with expiry control and a cost in SDG.")
    d.h2("4.2 Lot / batch and expiry tracking")
    d.p("Lot tracking records which **batch** each unit belongs to; expiry tracking records the "
        "**expiration date** of each batch. Together they give full traceability and drive FEFO.")
    d.ul(["On receipt the system asks for the **lot number** and **expiry date**.",
          "Stock can be traced by lot for recalls and audits.",
          "Near-expiry stock is highlighted and issued first under FEFO."])
    d.fig("lots_expiry.png", "Lot / batch records, each carrying an expiration date that drives FEFO.")
    d.h2("4.3 Checking stock on hand")
    d.p("Open a product and use the **On Hand** / **Forecasted** smart buttons, or review the "
        "on-hand column in the product list. **Inventory → Reporting** provides valuation and "
        "movement analysis.")

    d.h1("5. Warehouse Operations")
    d.fig("inventory.png", "The Inventory overview — each card is an operation type "
          "(receipts, deliveries, internal transfers) with its count of work to process.")
    d.h2("5.1 Receipts (goods in)")
    d.ol(["Receipts are generated from confirmed purchase orders.",
          "Open the receipt, enter the **lot number** and **expiry date** for each line, then "
          "**Validate**.",
          "Stock is added to the destination; putaway rules may route it (e.g. to Cold Storage)."])
    d.h2("5.2 Deliveries (goods out)")
    d.ol(["Deliveries are generated from confirmed sales orders.",
          "The system reserves stock — under **FEFO** it selects the soonest-to-expire lots.",
          "**Validate** the delivery to ship the goods and reduce stock."])
    d.h2("5.3 Internal transfers")
    d.p("Move stock between locations or warehouses — for example from the main store to "
        "**Cold Storage**. Use **Inventory → Operations → Transfers**, choose the source and "
        "destination, add the product and lot, and validate.")
    d.fig("transfers.png", "The Transfers list — receipts, deliveries and internal moves, "
          "each with its status (Done, Ready, Waiting).")
    d.h2("5.4 Reordering rules")
    d.p("Set minimum/maximum levels on critical items. When stock falls below the minimum the "
        "system proposes a purchase to top up to the maximum, preventing stock-outs.")

    d.h1("6. Procurement")
    d.h2("6.1 The purchasing flow")
    d.ol(["**Purchase → Orders → New** creates a Request for Quotation (RFQ).",
          "Select the supplier; add products and quantities.",
          "**Confirm Order** turns the RFQ into a Purchase Order and creates a **Receipt**.",
          "Receive the goods (entering lots and expiry).",
          "**Create Bill** generates the vendor bill; post it and register payment."])
    d.h2("6.2 Vendor price lists")
    d.p("On a product's **Purchase** tab, add one line per supplier with the price, currency, "
        "minimum quantity and lead time. Import suppliers can be priced in **USD**.")
    d.h2("6.3 Multi-currency purchasing")
    d.p("A purchase order can be issued in the supplier's currency (e.g. USD). The amount is "
        "converted to SDG using the exchange rate on the order/bill date, so the accounting "
        "value is always correct in the base currency.")
    d.tip("Bill control", "Set **Bill Control = Received quantities** so you only pay for what "
          "actually arrived — important for medical goods.")
    d.h2("6.4 Worked example — importing insulin in USD")
    d.p("This walkthrough uses the demo purchase order **P00002** placed with the import "
        "supplier **Gulf MedTrade FZE**, priced in **USD**.")
    d.ol(["In **Purchase → Orders**, the order **P00002** lists **50 × Insulin Vial 10ml** "
          "at **$14.00** and **10 × Digital Blood Pressure Monitor** at **$42.00**.",
          "Because Gulf MedTrade's purchase currency is **USD**, the order total is **$1,288.00**; "
          "the system shows the SDG value beneath it using the rate on the order date "
          "(1 USD = 700 SDG → about **901,600 SDG**).",
          "On **Confirm Order** a **Receipt** is created. Validating it (with lot **LOT-GULF-01** "
          "and the expiry date) adds the insulin to stock; the putaway rule routes it to "
          "**Cold Storage**.",
          "**Create Bill** generates the vendor bill in USD; it is valued in SDG when posted."])
    d.fig("purchase_usd.png", "Purchase order **P00002** from Gulf MedTrade — a USD order "
          "($1,288) automatically valued in the company's SDG accounts.")

    d.h1("7. Sales")
    d.h2("7.1 The selling flow")
    d.ol(["**Sales → Orders → New** creates a quotation.",
          "Select the customer; the customer's price list sets prices and currency.",
          "Add products and quantities and confirm the quotation.",
          "**Confirm** creates the Sales Order and a **Delivery**.",
          "Validate the delivery (FEFO selects the lots), then **Create Invoice** and post it."])
    d.h2("7.2 Price lists and currency")
    d.p("Each price list has a currency. Domestic sales use an **SDG** price list. The sales "
        "order's currency follows its price list, so assign the correct price list to each "
        "customer.")
    d.warn("Currency comes from the price list", "If invoices appear in the wrong currency, "
           "check the price list assigned to the customer — the order currency follows the "
           "price list, not the contact's country.")
    d.h2("7.3 Worked example — selling to Khartoum Teaching Hospital")
    d.p("This walkthrough follows the demo sales order **S00001** through to its posted invoice.")
    d.ol(["In **Sales → Orders**, order **S00001** sells to **Khartoum Teaching Hospital**: "
          "**50 × Paracetamol 500mg**, **20 × Surgical Gloves** and **30 × Syringe 5ml** — "
          "all priced from the **SDG** price list.",
          "Confirming the order creates a **Delivery**; under FEFO the system reserves the "
          "soonest-to-expire lots and is validated to ship the goods.",
          "**Create Invoice** then posts **INV/2026/00001** for **317,400 SDG** "
          "(276,000 net + 41,400 VAT at 15%)."])
    d.fig("sale_order.png", "Sales order **S00001** for Khartoum Teaching Hospital, priced in SDG.")

    d.h1("8. Accounting")
    d.h2("8.1 What is included")
    d.p("Community Odoo provides **Invoicing**; the bundled OCA / Odoo-Mates modules extend it "
        "into a full **Accounting** system.")
    d.table(["Capability", "Provided by"],
            [["Customer invoices, vendor bills, payments, taxes", "Invoicing (core)"],
             ["Full Accounting menu, asset, budget, fiscal year, recurring", "om_account_accountant suite"],
             ["Financial reports (Balance Sheet, P&L, ledgers, tax, aged)", "accounting_pdf_reports"],
             ["Cash Book / Day Book / Bank Book", "om_account_daily_reports"],
             ["Bank reconciliation (interactive matching)", "account_reconcile_oca"],
             ["Customer follow-up / dunning", "om_account_followup"]])
    d.fig("accounting.png", "The Accounting dashboard — one card per journal "
          "(Customer Invoices, Vendor Bills, Bank, Cash) with quick actions.")
    d.h2("8.2 Customer invoices and vendor bills")
    d.ul(["Customer invoices are normally created from sales orders and posted from Accounting.",
          "Vendor bills are created from purchase orders / receipts and posted before payment.",
          "Each posted document generates the matching journal entries automatically."])
    d.fig("customer_invoice.png", "Posted customer invoice **INV/2026/00001** in SDG, "
          "linked back to its sales order.")
    d.h2("8.3 Payments and the money safe")
    d.p("Register a payment from an invoice or bill and choose the journal — a **bank** account "
        "or a **cash safe**. Reconcile cash safes daily using the **Cash Book / Day Book**.")
    d.h2("8.4 Bank reconciliation")
    d.p("Use **Accounting → Bank** to match bank-statement lines against invoices and payments "
        "with the interactive reconciliation screen.")
    d.h2("8.5 Reports")
    d.ul(["**Balance Sheet** and **Profit & Loss** for financial position and performance.",
          "**General / Partner Ledger**, **Trial Balance**, **Aged Receivable / Payable**.",
          "**Tax Report** for filing; **Day/Cash/Bank Book** for daily cash control."])
    d.h2("8.6 Fiscal year and follow-up")
    d.ul(["Define the **Fiscal Year** and set **lock dates** after each period close.",
          "Configure **Follow-up levels** (e.g. 15 / 30 / 60 days) to chase overdue customers."])

    d.h1("9. Multi-Currency Operations")
    d.p("The system keeps accounts in **SDG** while letting you transact in **USD**.")
    d.ul(["**Base currency:** SDG — all ledgers and reports are in SDG.",
          "**Reference currency:** USD — used for import purchases and quotations.",
          "**Rates:** stored as dated lines; the rate on a document's date is used for conversion.",
          "**Foreign journals:** a USD bank or cash journal holds balances in USD and is "
          "revalued to SDG at period close."])
    d.table(["Document", "Currency source"],
            [["Purchase order", "Supplier's purchase currency / vendor price list"],
             ["Vendor bill", "Inherited from the purchase order"],
             ["Sales order", "Customer's price list"],
             ["Customer invoice", "Inherited from the sales order"]])

    d.h1("10. Medical-Supply Specifics")
    d.ul(["**Batch traceability** — every medicine is tracked by lot for recalls and audits.",
          "**Expiry control** — expiration dates are captured per lot and near-expiry stock is flagged.",
          "**FEFO** — deliveries automatically issue the earliest-expiring stock first.",
          "**Cold chain** — temperature-sensitive items (e.g. insulin) are stored in a Cold "
          "Storage location, with putaway rules routing them there on receipt.",
          "**Quarantine / damaged** — separate locations isolate non-sellable stock.",
          "**Imports** — purchases in USD are valued in SDG at the current rate; landed costs "
          "(freight, customs) can be folded into unit cost."])

    d.h1("11. Demo Data Reference")
    d.p("The demonstration database **erpmedsupply** is pre-loaded with realistic data so every "
        "feature can be explored immediately.")
    d.table(["Item", "Details"],
            [["Company", "Sudan MedSupply Co. (base currency SDG)"],
             ["Currencies", "SDG (base) + USD (reference), with 5 dated USD rates (600 → 700 SDG)"],
             ["Warehouses", "Khartoum Central (KRT), Port Sudan (PRT)"],
             ["Locations", "Cold Storage (2-8°C), Quarantine, Expired/Damaged"],
             ["Products", "10 medical items with lot & expiry tracking"],
             ["Customers", "Khartoum Teaching Hospital, Omdurman Family Clinic, Sudanese Red Crescent, Bahri Community Pharmacy"],
             ["Suppliers", "Nile Medical Supplies, Khartoum Pharma Imports, Gulf MedTrade (USD)"],
             ["Banks & safes", "2 bank journals (SDG/USD) + 2 cash safes (SDG/USD)"],
             ["Purchases", "3 purchase orders received (incl. one in USD); 1 vendor bill"],
             ["Sales", "2 delivered & invoiced orders (SDG) + 1 draft quotation"]])

    d.h1("12. Administration & Maintenance")
    d.h2("12.1 Users and access")
    d.p("Create users under **Settings → Users**. Assign application access rights so each role "
        "sees only the menus it needs.")
    d.h2("12.2 Backups")
    d.p("Back up the database regularly and keep off-site copies before upgrades.")
    d.h2("12.3 Rebuilding the demo")
    d.p("The demo database can be recreated from scratch using the project's seed script. See "
        "the project documentation for the exact commands.")
    d.warn("Production hardening", "Before real use: change all default passwords, disable the "
           "database manager, enable HTTPS, and restrict the database filter.")

    d.h1("13. Glossary")
    d.table(["Term", "Meaning"],
            [["SKU / Internal Reference", "Your unique product code."],
             ["Lot / Batch", "A group of units received together, tracked as one."],
             ["FEFO", "First Expiry First Out — issue soonest-to-expire stock first."],
             ["FIFO", "First In First Out — valuation using oldest cost first."],
             ["Putaway rule", "Automatic routing of received goods to a specific location."],
             ["Journal", "A book of account entries (Sales, Purchase, Bank, Cash, Misc)."],
             ["RFQ", "Request for Quotation — a draft purchase order."],
             ["Reconciliation", "Matching bank-statement lines to invoices and payments."],
             ["Base currency", "The currency the accounts are kept in (SDG)."]])

# ============================================================ ARABIC CONTENT
def fill_ar(d):
    d.h1("1. مقدمة")
    d.p("يصف هذا الدليل **نظام إدارة موارد المستلزمات الطبية** — وهو نظام أعمال متكامل لشركة "
        "تستورد وتخزّن وتوزّع المستلزمات الطبية والأدوية. يغطي النظام سلسلة العمليات الكاملة: "
        "إعداد البيانات الأساسية، والمخزون والمستودعات، والمشتريات، والمبيعات، والمحاسبة، مع دعم "
        "أصيل **لتعدد العملات** (الجنيه السوداني كعملة أساسية والدولار الأمريكي كعملة مرجعية).")
    d.p("النظام مبني على **أودو 18 المجتمعي** ويستخدم وحدات مجتمعية ومفتوحة المصدر فقط "
        "(OCA / Odoo-Mates) — دون الحاجة إلى أي تراخيص Enterprise مملوكة.")
    d.h2("1.1 لمن هذا الدليل")
    d.ul(["**المسؤولون** الذين يهيئون البيانات الأساسية — المستودعات والأصناف والعملاء والموردين "
          "والبنوك والعملات وأسعار الصرف.",
          "**موظفو المشتريات** الذين ينشئون أوامر الشراء ويستلمون البضائع.",
          "**موظفو المبيعات** الذين يعرضون الأسعار ويبيعون ويسلّمون للعملاء.",
          "**المحاسبون** الذين يديرون الفواتير والمدفوعات والتسوية البنكية والتقارير."])
    d.h2("1.2 لمحة سريعة عن القدرات")
    d.table(["المجال", "ما يمكنك عمله"],
            [["الإعدادات", "تعريف الشركة والعملات والأسعار والعملاء والموردين والبنوك وخزائن النقد والوحدات والمستودعات وفئات المنتجات."],
             ["المخزون", "تسجيل الأصناف مع تتبّع الدفعات وتواريخ الانتهاء؛ وإدارة الرصيد المتاح والتقييم."],
             ["المستودعات", "تشغيل عدة مستودعات ومواقع داخلية (مثل التخزين المبرّد والحجر)؛ والاستلام والتسليم والتحويلات الداخلية."],
             ["المشتريات", "طلب عرض سعر ← أمر شراء ← استلام ← فاتورة مورّد، مع قوائم أسعار الموردين والشراء بعملات متعددة."],
             ["المبيعات", "عرض سعر ← أمر بيع ← تسليم ← فاتورة عميل، مع قوائم أسعار لكل عملة."],
             ["المحاسبة", "فواتير العملاء وفواتير الموردين والمدفوعات والتسوية البنكية والتقارير المالية والسنة المالية والمطالبات."],
             ["تعدد العملات", "الجنيه السوداني أساسية والدولار مرجعية؛ مع سجل مؤرّخ لأسعار الصرف يسهل تحديثه."]])

    d.h1("2. نظرة عامة على النظام")
    d.h2("2.1 المكوّنات الأساسية")
    d.p("ينقسم النظام إلى تطبيقات متعاونة. يضيف كل تطبيق قوائمه ومستنداته وتقاريره، لكنها تتشارك "
        "نفس المنتجات وجهات الاتصال والمحاسبة.")
    d.table(["التطبيق", "الغرض"],
            [["المخزون", "الأصناف والرصيد والمستودعات والمواقع والدفعات وتواريخ الانتهاء والتحويلات."],
             ["المشتريات", "أوامر الموردين واستلام البضائع."],
             ["المبيعات", "عروض الأسعار وأوامر البيع والتسليمات."],
             ["الفوترة / المحاسبة", "الفواتير والمدفوعات ودفاتر اليومية والتقارير."],
             ["جهات الاتصال", "العملاء والموردون (يمكن أن تكون الجهة عميلًا ومورّدًا معًا)."],
             ["الإعدادات", "الشركة والعملات والمستخدمون ومفاتيح تفعيل الخصائص."]])
    d.h2("2.2 تسجيل الدخول")
    d.ol(["افتح رابط النظام في المتصفح (للنسخة التجريبية المحلية: **http://localhost:8069**).",
          "إذا ظهرت قائمة قواعد البيانات، اختر **erpmedsupply**.",
          "أدخل اسم المستخدم وكلمة المرور. المسؤول الافتراضي في النسخة التجريبية هو "
          "**admin / admin** — غيّر كلمة المرور قبل التشغيل الفعلي."])
    d.fig("login.png", "شاشة تسجيل الدخول — اختر قاعدة البيانات **erpmedsupply** وسجّل الدخول.")
    d.warn("غيّر بيانات الدخول الافتراضية", "يجب تغيير كلمة مرور المسؤول التجريبية (**admin**) "
           "وكلمة المرور الرئيسية لقاعدة البيانات قبل استخدام النظام ببيانات حقيقية.")
    d.h2("2.3 التنقّل في النظام")
    d.ul(["**شريط القوائم العلوي** ينقلك بين التطبيقات (المخزون، المشتريات، المبيعات، المحاسبة…).",
          "لكل تطبيق قائمة فرعية **الإعدادات** الخاصة به لمهام التهيئة.",
          "تتبع معظم المستندات دورة حياة **مسودة ← مؤكد ← منجز** تظهر عبر أزرار الحالة أعلى اليسار."])
    d.fig("apps_home.png", "الشاشة الرئيسية للتطبيقات، وتعرض تطبيقات الأعمال المثبّتة.")

    d.h1("3. الإعدادات (الإدارة المركزية)")
    d.p("الإعدادات هي قلب النظام. أكمل هذه الخطوات **بالترتيب** قبل تسجيل العمليات اليومية، لأن "
        "بعض الإعدادات (خاصة العملة الأساسية وشجرة الحسابات) يصعب تغييرها بعد وجود عمليات.")
    d.h2("3.1 الشركة")
    d.ol(["اذهب إلى **الإعدادات ← المستخدمون والشركات ← الشركات** وافتح شركتك.",
          "أدخل الاسم القانوني والعنوان والرقم الضريبي والشعار.",
          "اضبط **الدولة** على السودان."])
    d.h2("3.2 العملات وأسعار الصرف")
    d.p("يستخدم النظام **الجنيه السوداني (SDG)** كعملة أساسية للشركة و**الدولار الأمريكي (USD)** "
        "كعملة مرجعية للاستيراد وعروض الأسعار.")
    d.ol(["فعّل تعدد العملات: **الإعدادات ← المحاسبة ← العملات ← فعّل تعدد العملات**.",
          "فعّل العملات المستخدمة: **الإعدادات ← العملات**، فعّل **SDG** و**USD**.",
          "تأكد أن العملة الأساسية للشركة هي **SDG**."])
    d.h3("تسجيل سعر صرف جديد للدولار")
    d.p("تُخزَّن أسعار الصرف كـ**سجلات مؤرّخة**، فيُحفَظ التاريخ الكامل ويُقيَّم أي مستند سابق "
        "بالسعر الذي كان ساريًا في تاريخه.")
    d.ol(["افتح **المحاسبة ← الإعدادات ← العملات** واضغط **USD**.",
          "انتقل إلى تبويب **الأسعار** واضغط **إضافة سطر**.",
          "أدخل **التاريخ** والسعر — اكتب الرقم البديهي: **كم جنيهًا يساوي دولارًا واحدًا** (مثلًا 700).",
          "احفظ. ينطبق السعر الجديد على كل المستندات المؤرخة في ذلك اليوم أو بعده."])
    d.tip("تحديث سهل للأسعار", "بما أن كل تغيير سطر مؤرّخ جديد، فأنت تضيف سطرًا كلما تحرّك السعر "
          "— دون أن تطمس التاريخ. تستخدم التقارير والمستندات السعر الصحيح لتاريخها تلقائيًا.")
    d.fig("currency_rates.png", "سجل أسعار صرف الدولار المؤرّخ (1 دولار = 600 ← 700 جنيه). "
          "أدخل السعر في عمود **جنيه لكل وحدة**.")
    d.note("الأسعار التلقائية (اختياري)", "لا يجلب أودو المجتمعي الأسعار تلقائيًا. إن رغبت بتحديث "
           "مجدول، يمكن إضافة وحدة **currency_rate_update** من OCA لاحقًا؛ وإلا تُحدَّث الأسعار يدويًا.")
    d.h2("3.3 العملاء والموردون")
    d.p("العملاء والموردون كلاهما **جهات اتصال**. يمكن أن تكون الجهة الواحدة عميلًا ومورّدًا معًا "
        "— لا تنشئ نسخًا مكررة.")
    d.ol(["افتح **جهات الاتصال ← جديد**.",
          "اختر **شركة** أو **فرد** وأدخل الاسم والعنوان والهاتف.",
          "في تبويب **المبيعات والمشتريات** اضبط قائمة الأسعار وشروط الدفع و(للموردين الأجانب) "
          "عملة الشراء.",
          "في تبويب **المحاسبة** اضبط حسابات المدينين/الدائنين وأضف الحسابات البنكية."])
    d.table(["الحقل", "يُستخدم في"],
            [["عميل / مورّد", "يحدد ظهور الجهة في المبيعات أو المشتريات."],
             ["عملة الشراء", "العملة الافتراضية لأوامر هذا المورّد (مثل USD للاستيراد)."],
             ["قائمة الأسعار", "أسعار البيع والعملة التي تنطبق على هذا العميل."],
             ["الحسابات البنكية", "لازمة لتسجيل المدفوعات أو الدُفعات المجمّعة."]])
    d.fig("contacts.png", "دليل جهات الاتصال — العملاء والموردون في قائمة واحدة.")
    d.h2("3.4 البنوك وخزائن النقد")
    d.p("تُتابَع الحركات المالية عبر **دفاتر اليومية**. يحتاج **البنك** إلى سجل بنك ودفتر يومية "
        "بنكي؛ أما **خزينة النقد / الصندوق** فهي **دفتر يومية نقدي**.")
    d.ol(["أنشئ المؤسسة: **المحاسبة ← الإعدادات ← البنوك ← جديد** (الاسم، SWIFT/BIC).",
          "أنشئ دفتر يومية من نوع **بنك**. اضبط عملة للحساب بالعملة الأجنبية (مثل حساب بالدولار).",
          "أنشئ دفتر يومية **نقدي** لكل خزينة فعلية: النوع = **نقد** (مثل *الخزينة الرئيسية "
          "(SDG)*، *خزينة الدولار*)."])
    d.note("لماذا يمثّل الدفتر النقدي 'خزينة'", "لا يوجد كائن منفصل اسمه 'خزينة'. تُمثَّل كل خزينة "
           "أو صندوق فعلي بدفتر يومية **نقدي** خاص به، فيُتابَع رصيده وحركاته وتُسوّى باستقلالية.")
    d.h2("3.5 وحدات القياس")
    d.p("فعّل **وحدات القياس** للشراء والتخزين والبيع بوحدات مختلفة — مثلًا الشراء *بصندوق من "
        "100* والصرف *بالوحدة*. عرّف الوحدات ومعاملات التحويل من **المخزون ← الإعدادات ← وحدات القياس**.")
    d.h2("3.6 المستودعات والمواقع")
    d.ol(["اذهب إلى **المخزون ← الإعدادات ← المستودعات** وأنشئ مستودعًا لكل موقع فعلي "
          "(اسم + رمز مختصر، مثل *الخرطوم المركزي / KRT*).",
          "اضبط عدد خطوات الاستلام والتسليم لكل مستودع حسب الحاجة.",
          "أنشئ **مواقع** داخلية: **تخزين مبرّد (2-8°م)**، **حجر**، **منتهي/تالف**."])
    d.tip("سلسلة التبريد", "يُمثَّل التخزين المتحكَّم بحرارته كـ**موقع** داخلي. ويمكن لقاعدة "
          "**تخزين** أن توجّه منتجًا (مثل الإنسولين) تلقائيًا إلى التخزين المبرّد عند الاستلام.")
    d.fig("locations.png", "المواقع الداخلية ضمن مستودع الخرطوم المركزي — منها "
          "**التخزين المبرّد (2-8°م)** و**الحجر** و**منتهي/تالف**.")
    d.h2("3.7 فئات المنتجات")
    d.p("تتحكم فئات المنتجات في كيفية تسعير المخزون وتكلفته وطريقة صرفه.")
    d.table(["الإعداد", "القيمة الموصى بها", "الأثر"],
            [["طريقة التكلفة", "FIFO", "يُقيَّم المخزون بطريقة الوارد أولًا صادر أولًا."],
             ["تقييم المخزون", "يدوي / دوري*", "يتجنّب الحاجة لحسابات التقييم مسبقًا."],
             ["استراتيجية الصرف", "FEFO", "تُصرَف البضائع الأقرب انتهاءً أولًا."]])
    d.note("التقييم الآلي", "*تستخدم النسخة التجريبية التقييم الدوري لتُرحَّل حركات البضائع دون "
           "إعداد حسابات إضافية. لترحيل قيمة المخزون لحظيًا، اضبط حسابات الإدخال/الإخراج/التقييم "
           "ودفتر يومية المخزون للفئة، ثم بدّل التقييم إلى **آلي**.")

    d.h1("4. المخزون والأصناف")
    d.h2("4.1 تسجيل صنف")
    d.ol(["افتح **المخزون ← المنتجات ← المنتجات ← جديد**.",
          "أدخل الاسم و**المرجع الداخلي** (SKU) و**الباركود**.",
          "اضبط **نوع المنتج = بضاعة** وفعّل **تتبّع المخزون**.",
          "عيّن **فئة المنتج** (مثل الأدوية).",
          "اضبط **التكلفة** و**سعر البيع**.",
          "في تبويب **المخزون** اضبط **التتبّع = بالدفعات** وفعّل **تاريخ الانتهاء** للأدوية والمستهلكات."])
    d.fig("products_list.png", "كتالوج المنتجات الطبية مع الكميات المتاحة.")
    d.fig("product_insulin.png", "نموذج المنتج **قارورة إنسولين 10 مل** — متتبَّع بالدفعة "
          "مع ضبط الانتهاء وتكلفة بالجنيه.")
    d.h2("4.2 تتبّع الدفعات وتواريخ الانتهاء")
    d.p("يسجّل تتبّع الدفعات أي **دفعة** تنتمي إليها كل وحدة؛ ويسجّل تتبّع الانتهاء **تاريخ "
        "انتهاء** كل دفعة. ويوفّران معًا تتبّعًا كاملًا ويشغّلان آلية FEFO.")
    d.ul(["عند الاستلام يطلب النظام **رقم الدفعة** و**تاريخ الانتهاء**.",
          "يمكن تتبّع المخزون بالدفعة لأغراض السحب من السوق والتدقيق.",
          "يُبرَز المخزون القريب من الانتهاء ويُصرَف أولًا وفق FEFO."])
    d.fig("lots_expiry.png", "سجلات الدفعات، ويحمل كل منها تاريخ انتهاء يشغّل آلية FEFO.")
    d.h2("4.3 مراجعة الرصيد المتاح")
    d.p("افتح منتجًا واستخدم أزرار **المتاح** / **المتوقع**، أو راجع عمود الرصيد في قائمة "
        "المنتجات. توفّر **المخزون ← التقارير** تحليل التقييم والحركة.")

    d.h1("5. عمليات المستودع")
    d.fig("inventory.png", "نظرة عامة على المخزون — كل بطاقة نوع عملية (استلام، تسليم، "
          "تحويلات داخلية) مع عدد المهام المطلوب إنجازها.")
    d.h2("5.1 الاستلام (بضاعة واردة)")
    d.ol(["تُنشأ عمليات الاستلام من أوامر الشراء المؤكدة.",
          "افتح الاستلام، وأدخل **رقم الدفعة** و**تاريخ الانتهاء** لكل سطر، ثم **تحقّق**.",
          "يُضاف المخزون إلى الوجهة؛ وقد توجّهه قواعد التخزين (مثلًا إلى التخزين المبرّد)."])
    d.h2("5.2 التسليم (بضاعة صادرة)")
    d.ol(["تُنشأ عمليات التسليم من أوامر البيع المؤكدة.",
          "يحجز النظام المخزون — ووفق **FEFO** يختار الدفعات الأقرب انتهاءً.",
          "**تحقّق** من التسليم لشحن البضائع وخفض المخزون."])
    d.h2("5.3 التحويلات الداخلية")
    d.p("انقل المخزون بين المواقع أو المستودعات — مثلًا من المخزن الرئيسي إلى **التخزين "
        "المبرّد**. استخدم **المخزون ← العمليات ← التحويلات**، واختر المصدر والوجهة، وأضف المنتج "
        "والدفعة، ثم تحقّق.")
    d.fig("transfers.png", "قائمة التحويلات — الاستلام والتسليم والتحويلات الداخلية مع حالة كل منها.")
    d.h2("5.4 قواعد إعادة الطلب")
    d.p("اضبط حدودًا دنيا/عليا للأصناف الحرجة. عندما يهبط المخزون دون الحد الأدنى يقترح النظام "
        "أمر شراء لرفعه إلى الحد الأعلى، مانعًا نفاد المخزون.")

    d.h1("6. المشتريات")
    d.h2("6.1 مسار الشراء")
    d.ol(["**المشتريات ← الأوامر ← جديد** ينشئ طلب عرض سعر (RFQ).",
          "اختر المورّد؛ أضف المنتجات والكميات.",
          "**تأكيد الأمر** يحوّل الطلب إلى أمر شراء وينشئ **استلامًا**.",
          "استلم البضائع (مع إدخال الدفعات وتواريخ الانتهاء).",
          "**إنشاء فاتورة** يولّد فاتورة المورّد؛ رحّلها وسجّل الدفع."])
    d.h2("6.2 قوائم أسعار الموردين")
    d.p("في تبويب **المشتريات** للمنتج، أضف سطرًا لكل مورّد يتضمن السعر والعملة والكمية الدنيا "
        "ومدة التوريد. يمكن تسعير موردي الاستيراد **بالدولار**.")
    d.h2("6.3 الشراء بعملات متعددة")
    d.p("يمكن إصدار أمر الشراء بعملة المورّد (مثل الدولار). ويُحوَّل المبلغ إلى الجنيه السوداني "
        "بسعر الصرف في تاريخ الأمر/الفاتورة، فتكون القيمة المحاسبية صحيحة دائمًا بالعملة الأساسية.")
    d.tip("ضبط الفوترة", "اضبط **ضبط الفوترة = الكميات المستلمة** لتدفع فقط مقابل ما وصل فعلًا "
          "— وهو أمر مهم للبضائع الطبية.")
    d.h2("6.4 مثال تطبيقي — استيراد الإنسولين بالدولار")
    d.p("يستخدم هذا الشرح أمر الشراء التجريبي **P00002** الصادر للمورّد المستورد "
        "**Gulf MedTrade FZE** والمسعَّر **بالدولار**.")
    d.ol(["في **المشتريات ← الأوامر**، يتضمن الأمر **P00002**: **50 × قارورة إنسولين 10 مل** "
          "بسعر **14.00 دولار** و**10 × جهاز قياس ضغط رقمي** بسعر **42.00 دولار**.",
          "بما أن عملة شراء Gulf MedTrade هي **الدولار**، يبلغ إجمالي الأمر **1,288.00 دولار**؛ "
          "ويعرض النظام قيمته بالجنيه أسفله بسعر تاريخ الأمر (1 دولار = 700 جنيه ← نحو "
          "**901,600 جنيه**).",
          "عند **تأكيد الأمر** يُنشأ **استلام**. والتحقق منه (بالدفعة **LOT-GULF-01** وتاريخ "
          "الانتهاء) يضيف الإنسولين للمخزون؛ وتوجّهه قاعدة التخزين إلى **التخزين المبرّد**.",
          "**إنشاء فاتورة** يولّد فاتورة المورّد بالدولار؛ وتُقيَّم بالجنيه عند ترحيلها."])
    d.fig("purchase_usd.png", "أمر الشراء **P00002** من Gulf MedTrade — أمر بالدولار "
          "(1,288 دولار) يُقيَّم تلقائيًا في حسابات الشركة بالجنيه.")

    d.h1("7. المبيعات")
    d.h2("7.1 مسار البيع")
    d.ol(["**المبيعات ← الأوامر ← جديد** ينشئ عرض سعر.",
          "اختر العميل؛ تحدد قائمة أسعار العميل الأسعار والعملة.",
          "أضف المنتجات والكميات وأكّد العرض.",
          "**التأكيد** ينشئ أمر البيع و**تسليمًا**.",
          "تحقّق من التسليم (يختار FEFO الدفعات)، ثم **إنشاء فاتورة** ورحّلها."])
    d.h2("7.2 قوائم الأسعار والعملة")
    d.p("لكل قائمة أسعار عملة. تستخدم المبيعات المحلية قائمة أسعار **بالجنيه السوداني**. تتبع "
        "عملة أمر البيع قائمة أسعاره، فعيّن لكل عميل القائمة الصحيحة.")
    d.warn("العملة تأتي من قائمة الأسعار", "إن ظهرت الفواتير بعملة خاطئة، فتحقّق من قائمة الأسعار "
           "المعيّنة للعميل — تتبع عملة الأمر قائمة الأسعار لا دولة الجهة.")
    d.h2("7.3 مثال تطبيقي — البيع لمستشفى الخرطوم التعليمي")
    d.p("يتابع هذا الشرح أمر البيع التجريبي **S00001** حتى فاتورته المُرحَّلة.")
    d.ol(["في **المبيعات ← الأوامر**، يبيع الأمر **S00001** لـ**مستشفى الخرطوم التعليمي**: "
          "**50 × باراسيتامول 500مغ** و**20 × قفازات جراحية** و**30 × محقن 5مل** — وكلها "
          "مسعّرة من قائمة أسعار **الجنيه السوداني**.",
          "تأكيد الأمر ينشئ **تسليمًا**؛ ووفق FEFO يحجز النظام الدفعات الأقرب انتهاءً ثم "
          "يُتحقَّق منه لشحن البضائع.",
          "ثم **إنشاء فاتورة** يرحّل **INV/2026/00001** بمبلغ **317,400 جنيه** "
          "(276,000 صافٍ + 41,400 ضريبة بنسبة 15%)."])
    d.fig("sale_order.png", "أمر البيع **S00001** لمستشفى الخرطوم التعليمي، مسعّر بالجنيه السوداني.")

    d.h1("8. المحاسبة")
    d.h2("8.1 ما الذي يتضمنه")
    d.p("يوفّر أودو المجتمعي **الفوترة**؛ وتوسّعها وحدات OCA / Odoo-Mates المرفقة إلى نظام "
        "**محاسبة** كامل.")
    d.table(["القدرة", "مصدرها"],
            [["فواتير العملاء والموردين والمدفوعات والضرائب", "الفوترة (أساسي)"],
             ["قائمة محاسبة كاملة، الأصول، الموازنة، السنة المالية، المتكرر", "حزمة om_account_accountant"],
             ["تقارير مالية (الميزانية، الأرباح والخسائر، الدفاتر، الضريبة، الأعمار)", "accounting_pdf_reports"],
             ["دفتر النقد / اليومية / البنك", "om_account_daily_reports"],
             ["التسوية البنكية (المطابقة التفاعلية)", "account_reconcile_oca"],
             ["متابعة العملاء / المطالبات", "om_account_followup"]])
    d.fig("accounting.png", "لوحة المحاسبة — بطاقة لكل دفتر يومية (فواتير العملاء، فواتير "
          "الموردين، البنك، النقد) مع إجراءات سريعة.")
    d.h2("8.2 فواتير العملاء والموردين")
    d.ul(["تُنشأ فواتير العملاء عادةً من أوامر البيع وتُرحَّل من المحاسبة.",
          "تُنشأ فواتير الموردين من أوامر الشراء/الاستلام وتُرحَّل قبل الدفع.",
          "يولّد كل مستند مُرحَّل قيود اليومية المطابقة تلقائيًا."])
    d.fig("customer_invoice.png", "فاتورة عميل مُرحَّلة **INV/2026/00001** بالجنيه السوداني، "
          "مرتبطة بأمر بيعها.")
    d.h2("8.3 المدفوعات وخزينة النقد")
    d.p("سجّل دفعة من فاتورة واختر الدفتر — حساب **بنكي** أو **خزينة نقد**. سوِّ خزائن النقد "
        "يوميًا باستخدام **دفتر النقد / اليومية**.")
    d.h2("8.4 التسوية البنكية")
    d.p("استخدم **المحاسبة ← البنك** لمطابقة سطور كشف الحساب البنكي مع الفواتير والمدفوعات عبر "
        "شاشة التسوية التفاعلية.")
    d.h2("8.5 التقارير")
    d.ul(["**الميزانية العمومية** و**الأرباح والخسائر** للمركز والأداء الماليين.",
          "**دفتر الأستاذ العام / دفتر الشركاء**، **ميزان المراجعة**، **أعمار المدينين / الدائنين**.",
          "**تقرير الضريبة** للإقرار؛ و**دفتر اليومية/النقد/البنك** للرقابة النقدية اليومية."])
    d.h2("8.6 السنة المالية والمتابعة")
    d.ul(["عرّف **السنة المالية** واضبط **تواريخ الإقفال** بعد إغلاق كل فترة.",
          "هيّئ **مستويات المتابعة** (مثل 15 / 30 / 60 يومًا) لمطالبة العملاء المتأخرين."])

    d.h1("9. عمليات تعدّد العملات")
    d.p("يحتفظ النظام بالحسابات **بالجنيه السوداني** مع السماح لك بالتعامل **بالدولار**.")
    d.ul(["**العملة الأساسية:** SDG — كل الدفاتر والتقارير بالجنيه السوداني.",
          "**العملة المرجعية:** USD — تُستخدم لمشتريات الاستيراد وعروض الأسعار.",
          "**الأسعار:** تُخزَّن كسطور مؤرّخة؛ ويُستخدم سعر تاريخ المستند للتحويل.",
          "**الدفاتر الأجنبية:** يحتفظ دفتر بنك أو نقد بالدولار بأرصدته بالدولار ويُعاد تقييمه "
          "إلى الجنيه عند إقفال الفترة."])
    d.table(["المستند", "مصدر العملة"],
            [["أمر الشراء", "عملة شراء المورّد / قائمة أسعار المورّد"],
             ["فاتورة المورّد", "موروثة من أمر الشراء"],
             ["أمر البيع", "قائمة أسعار العميل"],
             ["فاتورة العميل", "موروثة من أمر البيع"]])

    d.h1("10. خصوصيات المستلزمات الطبية")
    d.ul(["**تتبّع الدفعات** — يُتتبَّع كل دواء بالدفعة لأغراض السحب والتدقيق.",
          "**ضبط الانتهاء** — تُسجَّل تواريخ الانتهاء لكل دفعة ويُبرَز المخزون القريب من الانتهاء.",
          "**FEFO** — تصرف التسليمات المخزون الأقرب انتهاءً أولًا تلقائيًا.",
          "**سلسلة التبريد** — تُخزَّن الأصناف الحساسة للحرارة (مثل الإنسولين) في موقع تخزين "
          "مبرّد، مع قواعد تخزين توجّهها إليه عند الاستلام.",
          "**الحجر / التالف** — مواقع منفصلة تعزل المخزون غير القابل للبيع.",
          "**الاستيراد** — تُقيَّم المشتريات بالدولار بالجنيه بالسعر الحالي؛ ويمكن دمج تكاليف "
          "الشحن والجمارك في تكلفة الوحدة."])

    d.h1("11. مرجع البيانات التجريبية")
    d.p("قاعدة البيانات التجريبية **erpmedsupply** محمّلة مسبقًا ببيانات واقعية لاستكشاف كل "
        "الخصائص فورًا.")
    d.table(["العنصر", "التفاصيل"],
            [["الشركة", "شركة السودان للمستلزمات الطبية (العملة الأساسية SDG)"],
             ["العملات", "SDG (أساسية) + USD (مرجعية)، مع 5 أسعار دولار مؤرّخة (600 ← 700 جنيه)"],
             ["المستودعات", "الخرطوم المركزي (KRT)، بورتسودان (PRT)"],
             ["المواقع", "تخزين مبرّد (2-8°م)، حجر، منتهي/تالف"],
             ["المنتجات", "10 أصناف طبية بتتبّع الدفعات والانتهاء"],
             ["العملاء", "مستشفى الخرطوم التعليمي، عيادة أم درمان، الهلال الأحمر السوداني، صيدلية بحري"],
             ["الموردون", "النيل للمستلزمات الطبية، الخرطوم للاستيراد الدوائي، غلف ميد تريد (بالدولار)"],
             ["البنوك والخزائن", "دفترا بنك (SDG/USD) + خزينتا نقد (SDG/USD)"],
             ["المشتريات", "3 أوامر شراء مستلمة (منها واحد بالدولار)؛ فاتورة مورّد واحدة"],
             ["المبيعات", "أمرا بيع مُسلَّمان ومُفوتران (بالجنيه) + عرض سعر مسودة"]])

    d.h1("12. الإدارة والصيانة")
    d.h2("12.1 المستخدمون والصلاحيات")
    d.p("أنشئ المستخدمين من **الإعدادات ← المستخدمون**. وعيّن صلاحيات الوصول للتطبيقات بحيث يرى "
        "كل دور القوائم التي يحتاجها فقط.")
    d.h2("12.2 النسخ الاحتياطي")
    d.p("خذ نسخًا احتياطية لقاعدة البيانات بانتظام واحتفظ بنسخ خارج الموقع قبل الترقيات.")
    d.h2("12.3 إعادة بناء النسخة التجريبية")
    d.p("يمكن إعادة إنشاء قاعدة البيانات التجريبية من الصفر باستخدام سكربت التهيئة الخاص "
        "بالمشروع. راجع وثائق المشروع للأوامر الدقيقة.")
    d.warn("تجهيز الإنتاج", "قبل الاستخدام الفعلي: غيّر كل كلمات المرور الافتراضية، وعطّل مدير "
           "قواعد البيانات، وفعّل HTTPS، وقيّد مرشّح قاعدة البيانات.")

    d.h1("13. مسرد المصطلحات")
    d.table(["المصطلح", "المعنى"],
            [["SKU / المرجع الداخلي", "رمز المنتج الفريد لديك."],
             ["الدفعة (Lot/Batch)", "مجموعة وحدات مستلمة معًا وتُتتبَّع كوحدة واحدة."],
             ["FEFO", "الأقرب انتهاءً صادر أولًا — يُصرَف المخزون الأقرب انتهاءً أولًا."],
             ["FIFO", "الوارد أولًا صادر أولًا — تقييم بأقدم تكلفة أولًا."],
             ["قاعدة التخزين", "توجيه آلي للبضائع المستلمة إلى موقع محدد."],
             ["دفتر اليومية", "سجل قيود محاسبية (مبيعات، مشتريات، بنك، نقد، متنوع)."],
             ["RFQ", "طلب عرض سعر — أمر شراء مسودة."],
             ["التسوية", "مطابقة سطور كشف الحساب البنكي مع الفواتير والمدفوعات."],
             ["العملة الأساسية", "العملة التي تُمسَك بها الحسابات (SDG)."]])

# ----------------------------------------------------------------------------
import re as _re
# paired single-asterisk italics: *text* (won't match lone footnote asterisks)
_ITALIC = _re.compile(r"\*([^*\s][^*]*?)\*")

def runs(text):
    """Yield (segment, bold, italic). Supports **bold** and paired *italic*."""
    out = []
    bold = False
    for seg in text.split("**"):
        if seg:
            pos = 0
            for m in _ITALIC.finditer(seg):
                if m.start() > pos:
                    out.append((seg[pos:m.start()], bold, False))
                out.append((m.group(1), bold, True))
                pos = m.end()
            if pos < len(seg):
                out.append((seg[pos:], bold, False))
        bold = not bold
    return out or [(text, False, False)]

# ============================================================ DOCX RENDERER
def build_docx(blocks, meta, path, rtl, font_name, img_dir):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)

    def style_heading(name, size, color):
        st = doc.styles[name]
        st.font.name = font_name
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
    style_heading("Heading 1", 18, BRAND_DARK)
    style_heading("Heading 2", 14, BRAND)
    style_heading("Heading 3", 12, ACCENT)

    def shade(cell, color):
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(sh)

    def set_rtl_par(par):
        pPr = par._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def cs_font(run):
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
        rfonts.set(qn("w:cs"), font_name)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        if rtl:
            rpr.append(OxmlElement("w:rtl"))

    def add_runs(par, text, color=None, bold_all=False):
        for seg, b, it in runs(text):
            r = par.add_run(seg)
            if b or bold_all:
                r.bold = True
            if it:
                r.italic = True
            if color:
                r.font.color.rgb = RGBColor.from_string(color)
            cs_font(r)
        if rtl:
            set_rtl_par(par)

    def heading(text, level):
        if level == 1:
            doc.add_page_break()
        par = doc.add_paragraph(style="Heading %d" % level)
        add_runs(par, text)
        return par

    # ---- Cover ----
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = band.rows[0].cells[0]
    shade(cell, BRAND)
    cell.paragraphs[0].text = ""
    cell.add_paragraph(""); cell.add_paragraph("")
    for txt, sz in [(meta["title"], 32), (meta["subtitle"], 20)]:
        pp = cell.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(txt); r.bold = True; r.font.size = Pt(sz)
        r.font.color.rgb = RGBColor.from_string("FFFFFF"); cs_font(r)
    cell.add_paragraph("")
    for txt, sz, col in [(meta["org"], 16, BRAND_DARK), (meta["platform"], 12, "333333"),
                         (meta["version"] + "  •  " + meta["date"], 12, "333333")]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(txt); r.font.size = Pt(sz)
        r.font.color.rgb = RGBColor.from_string(col)
        if txt == meta["org"]:
            r.bold = True
        cs_font(r)
    doc.add_page_break()

    # ---- TOC ----
    htoc = doc.add_paragraph()
    if rtl:
        set_rtl_par(htoc)
    r = htoc.add_run(meta["toc"]); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BRAND_DARK); cs_font(r)
    par = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    ph = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = meta["toc"]; ph.append(t); fld.append(ph); par._p.append(fld)
    doc.add_page_break()

    # ---- Footer page number ----
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    for el, attrs in [("w:fldChar", {"w:fldCharType": "begin"}),
                      ("w:instrText", None), ("w:fldChar", {"w:fldCharType": "end"})]:
        e = OxmlElement(el)
        if el == "w:instrText":
            e.set(qn("xml:space"), "preserve"); e.text = "PAGE"
        else:
            for k, v in attrs.items():
                e.set(qn(k), v)
        run._r.append(e)

    def callout(title, text, fill, bar):
        tbl = doc.add_table(rows=1, cols=1)
        if rtl:
            tbl._tbl.tblPr.append(OxmlElement("w:bidiVisual"))
        c = tbl.rows[0].cells[0]
        shade(c, fill)
        add_runs(c.paragraphs[0], title, color=bar, bold_all=True)
        add_runs(c.add_paragraph(), text)
        doc.add_paragraph("")

    def figure(filename, caption, n):
        src = os.path.join(img_dir, filename)
        if not os.path.exists(src):
            return
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run()
        run.add_picture(src, width=Inches(6.1))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if rtl:
            set_rtl_par(cap)
        label = "%s %d — %s" % (lab["fig"], n, caption)
        r = cap.add_run(label); r.italic = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("555555"); cs_font(r)
        doc.add_paragraph("")

    lab = meta["labels"]
    fig_no = 0
    for block in blocks:
        k = block[0]
        if k == "fig":
            fig_no += 1
            figure(block[1], block[2], fig_no)
        elif k == "h1":
            heading(block[1], 1)
        elif k == "h2":
            heading(block[1], 2)
        elif k == "h3":
            heading(block[1], 3)
        elif k == "p":
            add_runs(doc.add_paragraph(), block[1])
        elif k == "ul":
            for it in block[1]:
                add_runs(doc.add_paragraph(style="List Bullet"), it)
        elif k == "ol":
            for it in block[1]:
                add_runs(doc.add_paragraph(style="List Number"), it)
        elif k == "table":
            headers, rows_ = block[1], block[2]
            tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            if rtl:
                tbl._tbl.tblPr.append(OxmlElement("w:bidiVisual"))
            for i, htext in enumerate(headers):
                c = tbl.rows[0].cells[i]; shade(c, BRAND)
                add_runs(c.paragraphs[0], htext, color="FFFFFF", bold_all=True)
            for ri, row in enumerate(rows_):
                cells = tbl.add_row().cells
                for ci, val in enumerate(row):
                    if ri % 2 == 1:
                        shade(cells[ci], GREY)
                    add_runs(cells[ci].paragraphs[0], val)
            doc.add_paragraph("")
        elif k == "tip":
            callout(lab["tip"] + " — " + block[1], block[2], LIGHT, ACCENT)
        elif k == "warn":
            callout(lab["warn"] + " — " + block[1], block[2], "FBEAEA", "B00020")
        elif k == "note":
            callout(lab["note"] + " — " + block[1], block[2], GREY, BRAND_DARK)

    doc.save(path)

# ============================================================ HTML RENDERER
def slug(text):
    s = "".join(c.lower() if (c.isalnum()) else "-" for c in text)
    while "--" in s:
        s = s.replace("--", "-")
    return s.strip("-") or "sec"

def esc(t):
    out = []
    for s, b, it in runs(t):
        seg = _html.escape(s)
        if b:
            seg = "<strong>%s</strong>" % seg
        if it:
            seg = "<em>%s</em>" % seg
        out.append(seg)
    return "".join(out)

def build_html(blocks, meta, out_path, rtl, font_face, font_family, img_dir):
    import base64
    direction = "rtl" if rtl else "ltr"
    align = "right" if rtl else "left"
    side = "right" if rtl else "left"
    css = """
    %(face)s
    @page { margin: 18mm 16mm 16mm 16mm; }
    body { font-family:%(ff)s; color:#1a1a1a; font-size:11pt; line-height:1.6;
           margin:0; direction:%(dir)s; text-align:%(al)s; }
    h1 { color:#%(bd)s; font-size:21pt; border-bottom:3px solid #%(br)s; padding-bottom:6px;
         margin-top:0; page-break-before:always; }
    h2 { color:#%(br)s; font-size:15pt; margin-top:20px; }
    h3 { color:#%(ac)s; font-size:12.5pt; margin-top:14px; }
    p, li { font-size:11pt; }
    table { border-collapse:collapse; width:100%%; margin:12px 0; page-break-inside:avoid; }
    th { background:#%(br)s; color:#fff; text-align:%(al)s; padding:7px 9px; font-size:10.5pt; }
    td { border:1px solid #d6dde0; padding:6px 9px; vertical-align:top; font-size:10.5pt; }
    tr:nth-child(even) td { background:#%(gr)s; }
    .callout { border-radius:6px; padding:10px 14px; margin:12px 0; border-%(side)s:5px solid;
               page-break-inside:avoid; }
    .tip{background:#%(lt)s;border-color:#%(ac)s;} .warn{background:#FBEAEA;border-color:#B00020;}
    .note{background:#%(gr)s;border-color:#%(bd)s;}
    .callout .t{font-weight:bold;display:block;margin-bottom:3px;}
    .tip .t{color:#%(ac)s;} .warn .t{color:#B00020;} .note .t{color:#%(bd)s;}
    .cover{page-break-after:always;}
    .cover .band{background:#%(br)s;color:#fff;padding:150px 50px 80px;text-align:center;}
    .cover .band .t1{font-size:40pt;font-weight:bold;} .cover .band .t2{font-size:23pt;margin-top:10px;}
    .cover .meta{text-align:center;margin-top:80px;}
    .cover .meta .org{font-size:22pt;font-weight:bold;color:#%(bd)s;}
    .cover .meta .sub{font-size:13pt;margin-top:6px;color:#333;}
    .cover .meta .ver{font-size:13pt;margin-top:46px;color:#333;}
    .toc{page-break-after:always;}
    .toc h2{color:#%(bd)s;border-bottom:2px solid #%(br)s;padding-bottom:4px;}
    .toc ul{list-style:none;padding:0;} .toc li.l1{font-weight:bold;color:#%(bd)s;margin-top:8px;}
    .toc li.l2{padding-%(side)s:20px;color:#333;font-size:10.5pt;}
    .toc a{text-decoration:none;color:inherit;}
    figure{margin:14px 0;text-align:center;page-break-inside:avoid;}
    figure img{max-width:100%%;border:1px solid #cfd8dc;border-radius:5px;
               box-shadow:0 1px 4px rgba(0,0,0,0.12);}
    figcaption{font-size:9.5pt;color:#555;font-style:italic;margin-top:5px;}
    """ % {"face": font_face, "ff": font_family, "dir": direction, "al": align, "side": side,
           "bd": BRAND_DARK, "br": BRAND, "ac": ACCENT, "gr": GREY, "lt": LIGHT}

    toc = []
    for b in blocks:
        if b[0] in ("h1", "h2"):
            toc.append('<li class="%s"><a href="#%s">%s</a></li>'
                       % ("l1" if b[0] == "h1" else "l2", slug(b[1]), _html.escape(b[1])))

    parts = ['<!doctype html><html dir="%s"><head><meta charset="utf-8"><style>%s</style></head><body>'
             % (direction, css)]
    parts.append('<div class="cover"><div class="band"><div class="t1">%s</div>'
                 '<div class="t2">%s</div></div><div class="meta"><div class="org">%s</div>'
                 '<div class="sub">%s</div><div class="ver">%s &middot; %s</div></div></div>'
                 % (_html.escape(meta["title"]), _html.escape(meta["subtitle"]),
                    _html.escape(meta["org"]), _html.escape(meta["platform"]),
                    _html.escape(meta["version"]), _html.escape(meta["date"])))
    parts.append('<div class="toc"><h2>%s</h2><ul>%s</ul></div>' % (_html.escape(meta["toc"]), "".join(toc)))

    lab = meta["labels"]
    fig_no = 0
    for b in blocks:
        k = b[0]
        if k == "fig":
            src = os.path.join(img_dir, b[1])
            if not os.path.exists(src):
                continue
            fig_no += 1
            with open(src, "rb") as fh:
                data = base64.b64encode(fh.read()).decode("ascii")
            parts.append('<figure><img src="data:image/png;base64,%s"/>'
                         '<figcaption>%s %d — %s</figcaption></figure>'
                         % (data, _html.escape(lab["fig"]), fig_no, esc(b[2])))
        elif k in ("h1", "h2", "h3"):
            parts.append('<%s id="%s">%s</%s>' % (k, slug(b[1]), esc(b[1]), k))
        elif k == "p":
            parts.append("<p>%s</p>" % esc(b[1]))
        elif k == "ul":
            parts.append("<ul>%s</ul>" % "".join("<li>%s</li>" % esc(i) for i in b[1]))
        elif k == "ol":
            parts.append("<ol>%s</ol>" % "".join("<li>%s</li>" % esc(i) for i in b[1]))
        elif k == "table":
            th = "".join("<th>%s</th>" % esc(h) for h in b[1])
            trs = "".join("<tr>%s</tr>" % "".join("<td>%s</td>" % esc(c) for c in r) for r in b[2])
            parts.append("<table><tr>%s</tr>%s</table>" % (th, trs))
        elif k in ("tip", "warn", "note"):
            parts.append('<div class="callout %s"><span class="t">%s — %s</span>%s</div>'
                         % (k, lab[k], esc(b[1]), esc(b[2])))
    parts.append("</body></html>")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("".join(parts))

# ============================================================ MAIN
if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    base = "Medical-Supply_ERP_User_Manual"

    img_en = os.path.join(OUT_DIR, "img", "en")
    img_ar = os.path.join(OUT_DIR, "img", "ar")

    # Prefer the workflow-authored JSON chapters; fall back to fill_en/fill_ar.
    en_blocks = content_from_json("en")
    ar_blocks = content_from_json("ar")
    if not en_blocks:
        en = Content(); fill_en(en); en_blocks = en.C
    if not ar_blocks:
        ar = Content(); fill_ar(ar); ar_blocks = ar.C

    # English
    build_docx(en_blocks, META["en"], os.path.join(OUT_DIR, base + "_EN.docx"),
               rtl=False, font_name="Calibri", img_dir=img_en)
    build_html(en_blocks, META["en"], os.path.join(OUT_DIR, "_manual_EN.html"),
               rtl=False, font_face="", font_family="'Helvetica Neue',Arial,sans-serif",
               img_dir=img_en)

    # Arabic (RTL). Uses Alexandria — the same Arabic font applied to the system
    # (Spiffy theme). The PDF embeds Alexandria-{Regular,Bold}.ttf copied to the
    # container at /tmp/fonts/ by the build step.
    ar_face = ("@font-face{font-family:'Alexandria';font-weight:normal;"
               "src:url('file:///tmp/fonts/Alexandria-Regular.ttf');}"
               "@font-face{font-family:'Alexandria';font-weight:bold;"
               "src:url('file:///tmp/fonts/Alexandria-Bold.ttf');}")
    build_docx(ar_blocks, META["ar"], os.path.join(OUT_DIR, base + "_AR.docx"),
               rtl=True, font_name="Alexandria", img_dir=img_ar)
    build_html(ar_blocks, META["ar"], os.path.join(OUT_DIR, "_manual_AR.html"),
               rtl=True, font_face=ar_face, font_family="'Alexandria',sans-serif",
               img_dir=img_ar)

    print("EN blocks:", len(en_blocks), "| AR blocks:", len(ar_blocks))
    print("Outputs in", OUT_DIR)
